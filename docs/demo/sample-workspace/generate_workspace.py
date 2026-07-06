#!/usr/bin/env python3
"""
generate_workspace.py — builds the fictional advisor sample workspace used by
the Demo V1 run (docs/demo/sample-workspace/).

Purpose: produce believable client/household folders with REAL .docx and .pdf
files already on disk, so the demo can (a) look like a real practice on
screen and (b) exercise QA-92 ("Ask finds pre-existing files that were never
created inside the app").

Everything here is 100% fictional. The firm is "Beacon Ridge Wealth
Advisors" — never reference the product's real name or codename in generated
content.

Run from anywhere:
    python3 docs/demo/sample-workspace/generate_workspace.py

Requires: python-docx, fpdf2 (both already used/available in this repo's
Python environment).
"""
import os

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from fpdf.enums import XPos, YPos

HERE = os.path.dirname(os.path.abspath(__file__))
FIRM = "Beacon Ridge Wealth Advisors"
ADVISOR = "Sarah Kim, CFP"


# ---------------------------------------------------------------------------
# docx helpers
# ---------------------------------------------------------------------------

def new_doc(title, subtitle=None):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(11)
    doc.add_paragraph(f"{FIRM} — Prepared by {ADVISOR}")
    doc.add_paragraph("")
    return doc


def add_section(doc, heading, paragraphs):
    doc.add_heading(heading, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)


def add_table(doc, heading, header_row, rows):
    doc.add_heading(heading, level=1)
    table = doc.add_table(rows=1, cols=len(header_row))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, text in enumerate(header_row):
        hdr[i].text = text
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
    doc.add_paragraph("")


# ---------------------------------------------------------------------------
# pdf helper
# ---------------------------------------------------------------------------

_ASCII_MAP = str.maketrans({
    "—": "-",  # em dash
    "–": "-",  # en dash
    "‘": "'",
    "’": "'",
    "“": '"',
    "”": '"',
    "…": "...",
    "×": "x",
})


def _ascii(text):
    """Core PDF fonts (Helvetica) only support latin-1; flatten the smart
    punctuation used in the docx copy so fpdf2 doesn't choke."""
    return text.translate(_ASCII_MAP)


def write_pdf(path, title, blocks):
    """blocks: list of (heading_or_None, [lines]) tuples, rendered as plain
    left-aligned text so pdftotext extraction is trivial and reliable."""
    pdf = FPDF(format="Letter")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    def line(text, size, bold=False, height=6):
        pdf.set_font("Helvetica", "B" if bold else "", size)
        pdf.multi_cell(
            0, height, _ascii(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT
        )

    line(title, 16, bold=True, height=8)
    line(f"{FIRM}  |  Prepared by {ADVISOR}", 10)
    pdf.ln(4)

    for heading, lines in blocks:
        if heading:
            line(heading, 12, bold=True, height=7)
        for text in lines:
            line(text, 10)
        pdf.ln(3)

    pdf.output(path)


# ---------------------------------------------------------------------------
# Client 1 — The Hendersons
# ---------------------------------------------------------------------------

def build_hendersons(folder):
    os.makedirs(folder, exist_ok=True)

    # Financial Plan Summary
    doc = new_doc(
        "Financial Plan Summary — Robert & Linda Henderson",
        "Annual plan update — January 2026",
    )
    add_section(
        doc,
        "Household Overview",
        [
            "Robert Henderson (born March 14, 1953) and Linda Henderson (born "
            "August 2, 1955) are retired and living in Sarasota, FL. Robert "
            "retired from Gulf Coast Engineering in 2019; Linda retired from "
            "teaching in 2020.",
            "Robert turns 73 in March 2026, which triggers his first Required "
            "Minimum Distribution (RMD) from his Traditional IRA. Under the "
            "SECURE 2.0 rules, his first RMD must be taken by April 1, 2027.",
        ],
    )
    add_table(
        doc,
        "Current Assets (as of December 31, 2025)",
        ["Account", "Custodian", "Owner", "Balance"],
        [
            ["Traditional IRA", "Fidelity", "Robert", "$1,340,212"],
            ["Roth IRA", "Fidelity", "Linda", "$410,880"],
            ["Joint Brokerage", "Charles Schwab", "Joint", "$685,430"],
            ["Checking / Cash Reserve", "Schwab Bank", "Joint", "$62,000"],
        ],
    )
    add_section(
        doc,
        "RMD Projection",
        [
            "Using the IRS Uniform Lifetime Table divisor of 26.5 for age 73, "
            "Robert's first-year RMD is projected at approximately $50,566, "
            "based on the December 31, 2025 IRA balance of $1,340,212.",
            "Recommendation: withhold 15% federal / 4% Florida-equivalent "
            "reserve is not applicable (FL has no state income tax); "
            "quarterly estimated payments are not required given withholding "
            "elected directly from the distribution.",
        ],
    )
    add_section(
        doc,
        "Goal: Grandson Ethan's College Fund",
        [
            "Robert and Linda want to superfund a 529 plan for their grandson "
            "Ethan Henderson (age 9) with an $80,000 lump-sum contribution in "
            "2026, using the 5-year gift-tax averaging election "
            "($16,000/year annual exclusion × 5 years = $80,000 per "
            "grandparent-grandchild pair, filed on Form 709).",
            "Funding source: Charles Schwab joint brokerage account. "
            "Target 529 plan: Florida 529 Savings Plan, growth portfolio.",
        ],
    )
    add_section(
        doc,
        "Goal: Downsizing the Sarasota Home",
        [
            "The Hendersons plan to sell their current 4-bedroom home and "
            "move to a 2-bedroom condo in the same Sarasota community by "
            "mid-2027. Estimated net proceeds after a purchase: $350,000, "
            "which will be added to the joint brokerage account as a "
            "healthcare and long-term-care reserve.",
        ],
    )
    doc.save(os.path.join(folder, "Financial Plan Summary - Henderson.docx"))

    # Meeting Prep Notes
    doc = new_doc(
        "Meeting Prep Notes — Henderson Household",
        "Upcoming meeting: Thursday, January 15, 2026, 10:00 AM (office visit)",
    )
    add_section(
        doc,
        "Agenda",
        [
            "1. Confirm RMD election for Robert's Fidelity Traditional IRA — "
            "first distribution due by April 1, 2027; recommend taking it in "
            "December 2026 for cleaner tax-year matching.",
            "2. Walk through the $80,000 Form 709 superfunding gift for "
            "Ethan's 529 plan — confirm Robert and Linda both want to "
            "split-gift (each treated as gifting $40,000).",
            "3. Discuss Sarasota home downsizing timeline — they've met "
            "with a realtor (Coastal Realty Group) and are targeting a "
            "listing in Q2 2027.",
            "4. Review Q4 2025 Schwab brokerage statement together (attached "
            "separately) — YTD return 9.8%, slightly ahead of the IPS "
            "benchmark blend.",
        ],
    )
    add_section(
        doc,
        "Open Questions From Robert (via voicemail, Jan 8, 2026)",
        [
            "• “Does taking the RMD push us into a higher Medicare "
            "IRMAA bracket?” — need to model 2026 MAGI including the "
            "~$50,566 RMD before the meeting.",
            "• “Can we do the 529 gift in two contributions instead "
            "of one lump sum?” — yes, but the 5-year election still "
            "needs to be filed the year the first contribution is made.",
        ],
    )
    doc.save(os.path.join(folder, "Meeting Prep Notes - Henderson - 2026-01-15.docx"))

    # IPS Excerpt
    doc = new_doc(
        "Investment Policy Statement (Excerpt) — Henderson Household",
        "Effective date: February 1, 2025 — next scheduled review: February 2027",
    )
    add_section(
        doc,
        "Risk Tolerance & Time Horizon",
        [
            "Risk tolerance: Moderate. Time horizon: 15+ years (joint life "
            "expectancy), with an income-need horizon beginning immediately "
            "given retired status.",
            "Primary objective: capital preservation with modest growth to "
            "offset inflation, sufficient to support $9,500/month in "
            "portfolio-funded income alongside Social Security.",
        ],
    )
    add_table(
        doc,
        "Target Asset Allocation",
        ["Asset Class", "Target %", "Rebalance Band"],
        [
            ["U.S. Equity", "35%", "±5%"],
            ["International Equity", "20%", "±5%"],
            ["Fixed Income", "40%", "±5%"],
            ["Cash & Equivalents", "5%", "±2%"],
        ],
    )
    add_section(
        doc,
        "Constraints",
        [
            "Liquidity: maintain at least 12 months of anticipated "
            "distributions in cash/short-term instruments at all times.",
            "Tax considerations: prioritize the Traditional IRA for RMD "
            "withdrawals; use the taxable brokerage account for the 529 gift "
            "funding to preserve tax-deferred growth inside retirement "
            "accounts.",
        ],
    )
    doc.save(os.path.join(folder, "IPS Excerpt - Henderson.docx"))

    # PDF: statement summary
    write_pdf(
        os.path.join(folder, "Fidelity IRA Statement Summary - Q4 2025.pdf"),
        "Account Statement Summary — Fidelity Traditional IRA",
        [
            (None, ["Account holder: Robert Henderson", "Statement period: October 1, 2025 – December 31, 2025"]),
            (
                "Account Summary",
                [
                    "Beginning balance (10/1/2025): $1,247,905.10",
                    "Contributions: $0.00",
                    "Withdrawals: $0.00",
                    "Realized/unrealized gains: $92,307.34",
                    "Ending balance (12/31/2025): $1,340,212.44",
                    "Year-to-date return: 9.8%",
                ],
            ),
            (
                "Holdings (top positions)",
                [
                    "Vanguard Total Stock Market Index (VTSAX) — $482,340",
                    "Vanguard Total Bond Market Index (VBTLX) — $401,120",
                    "Vanguard Total International Stock Index (VTIAX) — $268,050",
                    "Money Market / Core Position — $188,702",
                ],
            ),
            (
                "Required Minimum Distribution Notice",
                [
                    "Robert Henderson turns 73 in 2026. Per IRS rules, the "
                    "first Required Minimum Distribution from this account "
                    "must be taken by April 1, 2027, based on the December "
                    "31, 2025 balance of $1,340,212.44 and the Uniform "
                    "Lifetime Table divisor of 26.5.",
                    "Projected first-year RMD: approximately $50,566.",
                ],
            ),
        ],
    )

    # PDF: signed agreement
    write_pdf(
        os.path.join(folder, "Beacon Ridge Advisory Agreement - Signed.pdf"),
        f"{FIRM} — Investment Advisory Services Agreement",
        [
            (None, ["Client: Robert & Linda Henderson", "Effective date: February 3, 2025"]),
            (
                "Scope of Services",
                [
                    f"{FIRM} will provide ongoing investment management, "
                    "financial planning, and tax-coordination advisory "
                    "services for the accounts listed in Schedule A, "
                    "including RMD planning and gifting strategy support.",
                ],
            ),
            (
                "Fee Schedule",
                [
                    "Annual advisory fee: 1.10% of assets under management, "
                    "billed quarterly in arrears, householded across all "
                    "Fidelity and Charles Schwab accounts listed in "
                    "Schedule A.",
                ],
            ),
            (
                "Signatures",
                [
                    "/s/ Robert Henderson        Date: February 3, 2025",
                    "/s/ Linda Henderson         Date: February 3, 2025",
                    f"/s/ {ADVISOR}, {FIRM}       Date: February 3, 2025",
                ],
            ),
        ],
    )


# ---------------------------------------------------------------------------
# Client 2 — Maria & Luis Alvarez
# ---------------------------------------------------------------------------

def build_alvarez(folder):
    os.makedirs(folder, exist_ok=True)

    doc = new_doc(
        "Financial Plan Summary — Maria & Luis Alvarez",
        "Annual plan update — February 2026",
    )
    add_section(
        doc,
        "Household Overview",
        [
            "Maria Alvarez (45) and Luis Alvarez (48) co-own Alvarez Family "
            "Taquerias, a 7-location restaurant group in the Austin, TX "
            "area. They have two children: Sofia (15) and Diego (12).",
            "A 2025 appraisal by Meridian Business Valuations placed the "
            "business at approximately $2,100,000 in enterprise value.",
        ],
    )
    add_table(
        doc,
        "Current Assets (as of December 31, 2025)",
        ["Account", "Custodian", "Owner", "Balance"],
        [
            ["Joint Brokerage", "Charles Schwab", "Joint", "$612,340"],
            ["SEP IRA", "Charles Schwab", "Luis", "$284,700"],
            ["Sofia's 529 (Texas Tuition Promise Fund)", "TX 529", "Sofia", "$58,000"],
            ["Diego's 529 (Texas Tuition Promise Fund)", "TX 529", "Diego", "$41,000"],
            ["Business Operating Reserve", "Chase Business", "Business", "$95,000"],
        ],
    )
    add_section(
        doc,
        "Goal: Business Exit Planning",
        [
            "Maria and Luis are targeting a sale of Alvarez Family Taquerias "
            "by 2031, once Sofia has finished college. They have engaged "
            "Meridian Business Valuations for annual appraisals to track "
            "enterprise value growth toward a $3,000,000+ exit target.",
        ],
    )
    add_section(
        doc,
        "Goal: Roth Conversion Strategy",
        [
            "Given a temporary dip in reportable business income in 2025 "
            "(remodel costs at two locations), we recommended converting "
            "$40,000 per year from Luis's SEP IRA to a Roth IRA in 2025, "
            "2026, and 2027, filling up the 22% federal bracket without "
            "spilling into the 24% bracket.",
        ],
    )
    add_section(
        doc,
        "Goal: Emergency Fund Build-Up",
        [
            "Current emergency reserve is $35,000; target is $75,000 "
            "(approximately 6 months of the couple's combined $12,500/month "
            "household draw from the business). Funding an additional "
            "$3,300/month from business distributions until the target is "
            "reached, expected by late 2026.",
        ],
    )
    doc.save(os.path.join(folder, "Financial Plan Summary - Alvarez.docx"))

    doc = new_doc(
        "Meeting Prep Notes — Alvarez Household",
        "Upcoming meeting: Wednesday, February 11, 2026, 4:00 PM (video call)",
    )
    add_section(
        doc,
        "Agenda",
        [
            "1. Q1 2026 estimated tax payments — confirm safe-harbor "
            "amount given the 2025 Roth conversion income.",
            "2. Tax-loss harvesting review on the Schwab joint brokerage "
            "account — two positions (an energy sector ETF and a "
            "regional bank fund) are down and eligible for harvesting "
            "before year-end planning.",
            "3. Business sale readiness checklist — walk through "
            "Meridian's 2025 appraisal letter and discuss what would move "
            "enterprise value toward the $3,000,000 target.",
            "4. Confirm 2026 Roth conversion amount ($40,000) and timing "
            "(recommend converting in two $20,000 tranches, April and "
            "October).",
        ],
    )
    add_section(
        doc,
        "Notes From Maria (email, Feb 3, 2026)",
        [
            "• “Diego wants to know if his 529 can be used for a "
            "trade school instead of a 4-year college — please check "
            "before the meeting.” (Answer: yes, 529 funds cover most "
            "accredited trade and vocational programs.)",
            "• “We got an unsolicited offer from a regional "
            "restaurant group for the business — want to discuss "
            "whether now is too early to sell.”",
        ],
    )
    doc.save(os.path.join(folder, "Meeting Prep Notes - Alvarez - 2026-02-11.docx"))

    doc = new_doc(
        "Investment Policy Statement (Excerpt) — Alvarez Household",
        "Effective date: March 15, 2024 — next scheduled review: March 2026",
    )
    add_section(
        doc,
        "Risk Tolerance & Time Horizon",
        [
            "Risk tolerance: Growth-oriented. Time horizon: 15+ years to "
            "the planned business exit and beyond, with two children's "
            "education funding needs beginning in 2029 (Sofia) and 2032 "
            "(Diego).",
        ],
    )
    add_table(
        doc,
        "Target Asset Allocation",
        ["Asset Class", "Target %", "Rebalance Band"],
        [
            ["U.S. Equity", "50%", "±7%"],
            ["International Equity", "25%", "±5%"],
            ["Fixed Income", "20%", "±5%"],
            ["Cash & Equivalents", "5%", "±2%"],
        ],
    )
    add_section(
        doc,
        "Constraints",
        [
            "Concentration risk: the household's net worth is heavily "
            "weighted toward the business itself; portfolio allocation "
            "intentionally favors liquid, diversified public equities to "
            "offset that concentration.",
        ],
    )
    doc.save(os.path.join(folder, "IPS Excerpt - Alvarez.docx"))

    write_pdf(
        os.path.join(folder, "Schwab Brokerage Statement Summary - Q4 2025.pdf"),
        "Account Statement Summary — Charles Schwab Joint Brokerage",
        [
            (None, ["Account holders: Maria & Luis Alvarez", "Statement period: October 1, 2025 – December 31, 2025"]),
            (
                "Account Summary",
                [
                    "Beginning balance (10/1/2025): $571,880.00",
                    "Contributions: $18,000.00",
                    "Withdrawals: $0.00",
                    "Realized/unrealized gains: $22,460.00",
                    "Ending balance (12/31/2025): $612,340.00",
                    "Year-to-date return: 11.2%",
                ],
            ),
            (
                "Positions Flagged for Tax-Loss Harvesting",
                [
                    "Energy Select Sector SPDR (XLE) — unrealized loss "
                    "of $4,120",
                    "Regional Bank ETF (KRE) — unrealized loss of $2,890",
                ],
            ),
        ],
    )

    write_pdf(
        os.path.join(folder, "Beacon Ridge Advisory Agreement - Signed.pdf"),
        f"{FIRM} — Investment Advisory Services Agreement",
        [
            (None, ["Client: Maria & Luis Alvarez", "Effective date: March 15, 2024"]),
            (
                "Scope of Services",
                [
                    f"{FIRM} will provide ongoing investment management, "
                    "business-owner financial planning, and Roth conversion "
                    "coordination for the accounts listed in Schedule A.",
                ],
            ),
            (
                "Fee Schedule",
                [
                    "Annual advisory fee: 1.25% of assets under management "
                    "(business-owner tier), billed quarterly in arrears.",
                ],
            ),
            (
                "Signatures",
                [
                    "/s/ Maria Alvarez           Date: March 15, 2024",
                    "/s/ Luis Alvarez            Date: March 15, 2024",
                    f"/s/ {ADVISOR}, {FIRM}       Date: March 15, 2024",
                ],
            ),
        ],
    )


# ---------------------------------------------------------------------------
# Client 3 — Dr. Priya Nair
# ---------------------------------------------------------------------------

def build_nair(folder):
    os.makedirs(folder, exist_ok=True)

    doc = new_doc(
        "Financial Plan Summary — Dr. Priya Nair",
        "Annual plan update — March 2026",
    )
    add_section(
        doc,
        "Client Overview",
        [
            "Dr. Priya Nair (37) is an interventional cardiologist at "
            "Cascade Regional Medical Center, a nonprofit hospital in the "
            "Seattle, WA area. She is single with no dependents. Current "
            "base income: $420,000/year.",
        ],
    )
    add_table(
        doc,
        "Current Assets (as of December 31, 2025)",
        ["Account", "Custodian", "Balance"],
        [
            ["403(b)", "Fidelity", "$215,000"],
            ["Taxable Brokerage", "Fidelity", "$95,000"],
            ["Roth IRA (backdoor contributions)", "Fidelity", "$41,300"],
            ["Cash Reserve", "Ally Bank", "$48,000"],
        ],
    )
    add_section(
        doc,
        "Goal: Student Loan Payoff via PSLF",
        [
            "Dr. Nair carries $178,400 in federal Direct and Grad PLUS "
            "student loans from medical school. Because Cascade Regional "
            "Medical Center is a qualifying nonprofit employer, she is "
            "pursuing Public Service Loan Forgiveness (PSLF).",
            "As of December 2025, she has made 84 of the required 120 "
            "qualifying payments, putting projected full forgiveness in "
            "2029. We recommend continuing minimum income-driven "
            "repayment plan payments rather than accelerating payoff.",
        ],
    )
    add_section(
        doc,
        "Goal: Backdoor Roth IRA",
        [
            "Dr. Nair's income exceeds the direct Roth IRA contribution "
            "limit, so she executes a backdoor Roth conversion each year: "
            "a $7,000 nondeductible contribution to a Traditional IRA "
            "(2026 limit), converted to Roth shortly after. She holds no "
            "other pre-tax IRA balances, so there is no pro-rata "
            "complication.",
        ],
    )
    add_section(
        doc,
        "Goal: Disability Insurance Review",
        [
            "Dr. Nair holds an own-occupation disability policy through "
            "Guardian with a $12,000/month benefit. Given her income "
            "growth, we recommend increasing coverage to $15,000/month at "
            "her next review, especially given the physical demands of "
            "interventional procedures.",
        ],
    )
    add_section(
        doc,
        "Goal: Home Purchase",
        [
            "Dr. Nair is saving toward a home purchase in the Seattle "
            "area, targeting homes in the $900,000–$1,100,000 range. "
            "Target down payment: $250,000 by summer 2027. Current "
            "dedicated home-fund savings: $61,000, saving an additional "
            "$5,500/month.",
        ],
    )
    doc.save(os.path.join(folder, "Financial Plan Summary - Nair.docx"))

    doc = new_doc(
        "Meeting Prep Notes — Dr. Priya Nair",
        "Upcoming meeting: Tuesday, March 3, 2026, 6:00 PM (video call, post-shift)",
    )
    add_section(
        doc,
        "Agenda",
        [
            "1. New attending contract offer from Cascade Regional — "
            "base salary increasing to $460,000 effective July 2026; "
            "confirm updated 403(b) contribution elections and whether "
            "PSLF nonprofit-employer status is unaffected (it is not).",
            "2. Annual PSLF employment certification form — due by "
            "April 30, 2026; confirm HR has processed last year's form.",
            "3. Disability insurance increase to $15,000/month — "
            "review Guardian's updated quote before the meeting.",
            "4. Home purchase savings pace check-in — on track for "
            "the $250,000 target by summer 2027 at current savings rate.",
        ],
    )
    add_section(
        doc,
        "Notes From Dr. Nair (text message, Feb 24, 2026)",
        [
            "• “If I take the new contract, does the higher "
            "income mess up my backdoor Roth?” (Answer: no — the "
            "backdoor Roth strategy is income-limit-proof by design.)",
            "• “Should I put anything extra toward the student "
            "loans now that I'm this close to forgiveness?” (Answer: "
            "no — extra payments would reduce the forgiven amount "
            "with no benefit.)",
        ],
    )
    doc.save(os.path.join(folder, "Meeting Prep Notes - Nair - 2026-03-03.docx"))

    doc = new_doc(
        "Investment Policy Statement (Excerpt) — Dr. Priya Nair",
        "Effective date: June 20, 2025 — next scheduled review: June 2027",
    )
    add_section(
        doc,
        "Risk Tolerance & Time Horizon",
        [
            "Risk tolerance: Aggressive growth. Time horizon: 25+ years to "
            "traditional retirement age. High and growing earned income "
            "provides significant capacity to bear market volatility.",
        ],
    )
    add_table(
        doc,
        "Target Asset Allocation",
        ["Asset Class", "Target %", "Rebalance Band"],
        [
            ["U.S. Equity", "55%", "±7%"],
            ["International Equity", "30%", "±7%"],
            ["Fixed Income", "10%", "±3%"],
            ["Cash & Equivalents", "5%", "±2%"],
        ],
    )
    add_section(
        doc,
        "Constraints",
        [
            "Home purchase carve-out: the dedicated home-fund savings "
            "(currently $61,000) are held outside this allocation, in "
            "high-yield savings, since the funds are needed within 24 "
            "months.",
        ],
    )
    doc.save(os.path.join(folder, "IPS Excerpt - Nair.docx"))

    write_pdf(
        os.path.join(folder, "Fidelity Statement Summary - Q4 2025.pdf"),
        "Account Statement Summary — Fidelity 403(b) & Brokerage",
        [
            (None, ["Account holder: Dr. Priya Nair", "Statement period: October 1, 2025 – December 31, 2025"]),
            (
                "403(b) Summary",
                [
                    "Beginning balance (10/1/2025): $198,400.00",
                    "Contributions: $9,200.00",
                    "Realized/unrealized gains: $7,400.00",
                    "Ending balance (12/31/2025): $215,000.00",
                ],
            ),
            (
                "Taxable Brokerage Summary",
                [
                    "Beginning balance (10/1/2025): $84,150.00",
                    "Contributions: $6,500.00",
                    "Realized/unrealized gains: $4,350.00",
                    "Ending balance (12/31/2025): $95,000.00",
                ],
            ),
        ],
    )

    write_pdf(
        os.path.join(folder, "Beacon Ridge Advisory Agreement - Signed.pdf"),
        f"{FIRM} — Investment Advisory Services Agreement",
        [
            (None, ["Client: Dr. Priya Nair", "Effective date: June 20, 2025"]),
            (
                "Scope of Services",
                [
                    f"{FIRM} will provide ongoing investment management, "
                    "student loan / PSLF coordination guidance, and "
                    "insurance review services for the accounts listed in "
                    "Schedule A.",
                ],
            ),
            (
                "Fee Schedule",
                [
                    "Annual advisory fee: 1.00% of assets under management "
                    "(young professional tier), billed quarterly in "
                    "arrears.",
                ],
            ),
            (
                "Signatures",
                [
                    "/s/ Priya Nair, M.D.        Date: June 20, 2025",
                    f"/s/ {ADVISOR}, {FIRM}       Date: June 20, 2025",
                ],
            ),
        ],
    )


def main():
    build_hendersons(os.path.join(HERE, "The Hendersons"))
    build_alvarez(os.path.join(HERE, "Maria & Luis Alvarez"))
    build_nair(os.path.join(HERE, "Dr. Priya Nair"))
    print("Sample workspace generated under:", HERE)


if __name__ == "__main__":
    main()
