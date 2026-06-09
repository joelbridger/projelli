#!/usr/bin/env python3
"""Independent (non-Keepance) validation of a produced .docx.

Uses `python-docx` (which itself uses lxml) as a SECOND opinion that our
in-house OOXML writer produces a file a real third-party Word-format reader
can open and understand. python-docx has no first-class track-changes API, so
we (a) confirm the package opens via Document(), then (b) walk the underlying
lxml tree for w:ins / w:del / commentRangeStart / commentReference and the
comments part, asserting the revision + comment markup is present and readable.

This is NOT a substitute for opening in real Microsoft Word — it cannot verify
Word's accept/reject UI. It proves the package is well-formed enough for an
independent OOXML consumer to parse the revisions and comments.

Usage: python3 independent_validate.py <file.docx> [<file2.docx> ...]
Exit code 0 = all files passed, 1 = a failure.
"""
import sys
import zipfile

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def check(path: str) -> list[str]:
    errors: list[str] = []

    # (0) It must be a valid zip with the required parts.
    try:
        with zipfile.ZipFile(path) as z:
            names = set(z.namelist())
    except Exception as e:  # noqa: BLE001
        return [f"not a valid zip: {e}"]
    for req in ("[Content_Types].xml", "_rels/.rels", "word/document.xml"):
        if req not in names:
            errors.append(f"missing part: {req}")

    # (1) python-docx — the independent reader — must open it.
    try:
        import docx  # python-docx
    except ImportError:
        errors.append("python-docx not installed; skipping independent open")
        return errors

    try:
        d = docx.Document(path)
    except Exception as e:  # noqa: BLE001
        return errors + [f"python-docx failed to OPEN the document: {e}"]

    # (2) The visible text python-docx reconstructs (it concatenates w:t runs,
    #     which for our model means insertions show up but deletions, stored as
    #     w:delText, are correctly NOT in the body text — exactly Word's
    #     "show final" behavior).
    body_text = "\n".join(p.text for p in d.paragraphs)
    if not body_text.strip():
        errors.append("python-docx read an EMPTY body (runs unreadable)")

    # (3) Walk the raw element tree for revision + comment markup.
    body = d.element.body
    ins = body.findall(f".//{W}ins")
    dele = body.findall(f".//{W}del")
    cr_start = body.findall(f".//{W}commentRangeStart")
    cr_ref = body.findall(f".//{W}commentReference")

    for el in ins + dele:
        for a in ("id", "author", "date"):
            if f"{W}{a}" not in el.attrib:
                errors.append(f"<{el.tag}> missing w:{a} attribute")

    # delText must live under w:del and carry the deleted text.
    for de in dele:
        dts = de.findall(f".//{W}delText")
        if not dts:
            errors.append("<w:del> has no <w:delText> child")
        elif not any((t.text or "") for t in dts):
            errors.append("<w:delText> is empty")

    # (4) comments.xml, if referenced, must exist and resolve.
    referenced_ids = {el.get(f"{W}id") for el in cr_ref}
    if referenced_ids:
        if "word/comments.xml" not in names:
            errors.append("comments referenced but word/comments.xml missing")
        else:
            from lxml import etree

            with zipfile.ZipFile(path) as z:
                ctree = etree.fromstring(z.read("word/comments.xml"))
            comment_ids = {c.get(f"{W}id") for c in ctree.findall(f"{W}comment")}
            for rid in referenced_ids:
                if rid not in comment_ids:
                    errors.append(f"commentReference id {rid} not found in comments.xml")

    # Report what we saw (stdout, for the human reading the spike output).
    print(f"  [{path}]")
    print(f"    python-docx opened OK; body text = {body_text!r}")
    print(
        f"    revisions: {len(ins)} <w:ins>, {len(dele)} <w:del>; "
        f"comments: {len(cr_start)} range(s), {len(cr_ref)} ref(s)"
    )
    authors = sorted(
        {el.get(f"{W}author") for el in ins + dele if el.get(f"{W}author")}
    )
    print(f"    revision authors: {authors}")
    return errors


def main() -> int:
    files = sys.argv[1:]
    if not files:
        print("usage: independent_validate.py <file.docx> [...]", file=sys.stderr)
        return 2
    failed = False
    for f in files:
        errs = check(f)
        if errs:
            failed = True
            print(f"FAIL {f}:")
            for e in errs:
                print(f"    - {e}")
        else:
            print(f"PASS {f}")
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
